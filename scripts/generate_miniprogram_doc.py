#!/usr/bin/env python3
"""
AI 教学辅助系统 — 微信小程序移动端 UI 设计方案 生成器
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime
import os

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs')
OUTPUT_FILE = os.path.join(OUTPUT_DIR, '微信小程序移动端UI设计方案.docx')

FONT_CN_HEADING = 'SimHei'
FONT_CN_BODY = 'SimSun'
FONT_EN = 'Times New Roman'
C_PRIMARY = RGBColor(0x1A, 0x56, 0xDB)
C_BODY = RGBColor(0x18, 0x20, 0x30)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_ACCENT = RGBColor(0xE6, 0x7E, 0x22)
BG_HEADER = '1A56DB'
BG_ALT = 'F0F4FF'
BG_COVER_TOP = '1A56DB'

def shade(cell, color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def add_run(p, text, bold=False, size=Pt(10.5), font_cn=FONT_CN_BODY, font_en=FONT_EN, color=C_BODY, italic=False):
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = size; r.font.name = font_en; r.font.color.rgb = color
    rPr = r._element.get_or_add_rPr()
    rPr.insert(0, parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_cn}" w:ascii="{font_en}" w:hAnsi="{font_en}"/>'))
    return r

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = FONT_EN; r.font.color.rgb = C_PRIMARY
        rPr = r._element.get_or_add_rPr()
        rPr.insert(0, parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_CN_HEADING}" w:ascii="{FONT_EN}" w:hAnsi="{FONT_EN}"/>'))
    return h

def body(doc, text, indent=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.space_before = Pt(2); pf.space_after = Pt(2)
    if indent: pf.first_line_indent = Pt(21)
    add_run(p, text)
    return p

def flow(doc, lines):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.left_indent = Cm(0.8); pf.space_before = Pt(2); pf.space_after = Pt(2)
    for i, line in enumerate(lines):
        if i > 0: p.add_run('\n')
        add_run(p, line, font_cn='Microsoft YaHei', size=Pt(8.5), color=RGBColor(0x50, 0x50, 0x50))
    return p

def table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=Pt(9), font_cn=FONT_CN_HEADING, color=C_WHITE)
        shade(c, BG_HEADER)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]; c.text = ''
            add_run(c.paragraphs[0], str(val), size=Pt(9))
            if ri % 2 == 1: shade(c, BG_ALT)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells): row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return tbl

def page_break(doc):
    doc.add_page_break()

# ============ 创建文档 ============
doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0); section.page_height = Cm(29.7)
section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.0); section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = FONT_EN; style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN_BODY)

# ============ 封面 ============
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'AI 教学辅助系统', bold=True, size=Pt(28), font_cn=FONT_CN_HEADING, color=C_PRIMARY)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '微信小程序移动端 UI 设计方案', bold=True, size=Pt(20), font_cn=FONT_CN_HEADING, color=C_ACCENT)

doc.add_paragraph(); doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '基于现有 Web 系统页面/接口/数据结构的移动端适配方案', size=Pt(11), color=RGBColor(0x50, 0x60, 0x70))

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, f'版本：V1.0 | 日期：{datetime.date.today().strftime("%Y-%m-%d")}', size=Pt(10), color=RGBColor(0x70, 0x80, 0x90))

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '设计原则：功能不变 · 字段不变 · 接口不变 · 流程不变', size=Pt(9), italic=True, color=RGBColor(0x90, 0x90, 0x90))

page_break(doc)

# ============ 目录 ============
heading(doc, '目录', level=1)
p = doc.add_paragraph()
add_run(p, '（请在 Word 中右键此处 → 更新域 → 更新整个目录）', size=Pt(9), italic=True, color=RGBColor(0x80, 0x90, 0xA0))

# TOC 域
para = doc.add_paragraph()
r = para.add_run()
r._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
r2 = para.add_run()
r2._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'))
r3 = para.add_run()
r3._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
r4 = para.add_run('（请右键更新目录）')
r4._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

page_break(doc)

# ============ 0. 设计前提 ============
heading(doc, '0. 设计前提', level=1)
body(doc, '当前 Web 系统核心页面包括 /login、/home、/workspace、/personal-lessons、/group-lessons、/plans、/reflections、/teacher/[id] 以及管理后台页面；其中 /workspace 是 AI 对话备课核心页，四个内容库页面复用 ContentLibraryPage 组件，仅通过 kind 区分个人备课、集体备课、计划总结、教学反思。')

heading(doc, '小程序设计原则', level=2)
table(doc, ['原则', '说明'],
    [['功能不变', '所有业务功能与 Web 端一致'],
     ['字段不变', '所有数据字段与 Web 端一致'],
     ['接口不变', '所有 API 接口与 Web 端一致'],
     ['流程不变', '所有业务流程与 Web 端一致'],
     ['Web表格→移动端卡片', '表格改为卡片列表'],
     ['Web三栏→页面+抽屉+底部栏', '三栏布局拆分为主页面+抽屉+底部操作栏'],
     ['AI工作流不变', '输入→AI处理→结构化预览→用户确认→保存入库']],
    col_widths=[3.5, 8.5])

page_break(doc)

# ============ 1. 小程序页面结构设计 ============
heading(doc, '1. 小程序页面结构设计', level=1)

heading(doc, '1.1 底部 Tab 结构', level=2)
body(doc, '小程序底部固定 4 个 Tab：')
flow(doc, ['首页    AI工作台    空间    我的'])

heading(doc, '1.2 Web 页面到小程序页面映射', level=2)
table(doc, ['Web 页面', 'Web 路由', '小程序页面', '说明'],
    [['登录页', '/login', '/pages/login/index', '登录、验证码、首次改密'],
     ['首页教师名录', '/home', '/pages/home/index', '教师名录、分组、搜索'],
     ['AI 工作台', '/workspace', '/pages/workspace/index', 'AI 对话备课核心页'],
     ['个人备课库', '/personal-lessons', '/pages/space/personal-lessons/index', '内容卡片列表'],
     ['集体备课库', '/group-lessons', '/pages/space/group-lessons/index', '内容卡片列表'],
     ['计划与总结库', '/plans', '/pages/space/plans/index', '内容卡片列表'],
     ['教学反思库', '/reflections', '/pages/space/reflections/index', '内容卡片列表'],
     ['教师个人空间', '/teacher/[id]', '/pages/teacher/detail/index', '教师统计+内容列表'],
     ['内容详情抽屉', 'Drawer', '/pages/content/detail/index', '移动端改为独立详情页'],
     ['AI预览卡编辑', '内联编辑', '底部弹层', '不改接口，只改展示'],
     ['管理后台', '/admin/*', '/pages/mine/admin-entry/index', '保留入口，跳转H5或分包']],
    col_widths=[2.5, 2.5, 3.5, 3.5])

page_break(doc)

# ============ 2. 页面逐个 UI 方案 ============
heading(doc, '2. 页面逐个 UI 方案', level=1)

# --- 2.1 登录页 ---
heading(doc, '2.1 登录页', level=2)
body(doc, '对应 Web 功能：Web 登录页包含学校品牌信息、手机号、密码、验证码、登录按钮，以及首次登录强制改密弹窗。')

heading(doc, '页面布局', level=3)
flow(doc, ['顶部：学校 Logo + 学校名称', '中部：登录卡片', '底部：系统说明 / 协议文字', '弹层：首次登录修改密码'])

heading(doc, '组件结构', level=3)
table(doc, ['区域', '组件'],
    [['Header', 'SchoolLogo、SchoolName'],
     ['Content', 'LoginCard、MobileInput、PasswordInput、CaptchaInput'],
     ['Modal', 'ChangePasswordModal'],
     ['Feedback', 'ErrorToast、LoadingButton']],
    col_widths=[3, 9])

heading(doc, '状态展示', level=3)
table(doc, ['状态', '移动端表现'],
    [['验证码加载中', '骨架占位'],
     ['登录中', '按钮 loading'],
     ['登录失败', '顶部轻提示 + 自动刷新验证码'],
     ['首次改密', '半屏弹窗'],
     ['改密失败', '表单错误提示'],
     ['登录成功', '跳转首页']],
    col_widths=[3, 9])

heading(doc, '移动端交互', level=3)
body(doc, '点击验证码图片刷新；点击眼睛图标切换密码显示；改密弹窗不可随意关闭，必须完成后进入系统。')

# --- 2.2 首页 ---
page_break(doc)
heading(doc, '2.2 首页', level=2)
body(doc, '对应 Web 功能：Web 首页是教师名录，展示全校教师列表，按部门或首页分组展示，支持搜索。')

heading(doc, '页面布局', level=3)
flow(doc, ['Header：学校名称 / 问候语 / 搜索入口', 'Content：搜索框 → 分组横向筛选 → 教师分组卡片 → 教师卡片网格', 'Tab：底部导航'])

heading(doc, '组件结构', level=3)
table(doc, ['区域', '组件'],
    [['Header', 'HomeHeader、SearchBar'],
     ['Content', 'GroupTabs、TeacherGroupCard、TeacherMiniCard'],
     ['Empty', 'EmptyState'],
     ['Tab', 'BottomTabBar']],
    col_widths=[3, 9])

heading(doc, '教师卡片展示', level=3)
flow(doc, ['头像 / 默认头像', '教师姓名', '所属分组', '个人备课数', '教学反思数'])

heading(doc, '状态展示', level=3)
table(doc, ['状态', '移动端表现'],
    [['加载中', '教师卡片骨架屏'],
     ['空数据', '\u201c暂无教师信息\u201d'],
     ['搜索无结果', '\u201c未找到相关教师\u201d'],
     ['接口错误', '错误卡片 + 重试按钮']],
    col_widths=[3, 9])

heading(doc, '移动端交互', level=3)
body(doc, '分组横向滑动；教师卡片点击进入详情；搜索框吸顶；下拉刷新首页数据。')

# --- 2.3 AI 工作台 ---
page_break(doc)
heading(doc, '2.3 AI 工作台', level=2)
body(doc, '对应 Web 功能：三栏布局（左会话列表 / 中聊天区 / 右工作列表）。小程序拆分为：主页面AI聊天 + 顶部抽屉会话列表 + 底部抽屉最近内容 + 预览卡出现在聊天流 + 详情页跳转。')

heading(doc, '页面布局', level=3)
flow(doc, [
    'Header：返回 / 会话标题 / 会话列表按钮',
    'ModeBar：自动识别 / 普通聊天 / 个人备课 / 集体备课 / 计划 / 总结 / 反思',
    'Content：聊天消息流',
    'BottomInput：+附件 / 输入框 / 发送',
    'Floating：最近保存内容入口',
])

heading(doc, '组件结构', level=3)
table(doc, ['区域', '组件'],
    [['Header', 'WorkspaceHeader'],
     ['Mode', 'ModeSelector'],
     ['Chat', 'MessageList、UserBubble、AIBubble、AIPreviewCard'],
     ['Input', 'ChatInputBar、AttachmentMenu'],
     ['Drawer', 'ConversationDrawer、RecentContentDrawer'],
     ['Modal', 'ConfirmSaveModal、CancelConfirmModal']],
    col_widths=[3, 9])

heading(doc, 'AI 结构化预览卡展示', level=3)
flow(doc, [
    '\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510',
    '\u2502 AI\u8bc6\u522b\u7ed3\u679c          \u2502',
    '\u2502 \u7c7b\u578b\uff1a\u4e2a\u4eba\u5907\u8bfe       \u2502',
    '\u2502 \u6807\u9898\uff1a\u300a\u8377\u5858\u6708\u8272\u300b   \u2502',
    '\u2502 \u5b66\u79d1\uff1a\u8bed\u6587           \u2502',
    '\u2502 \u5e74\u7ea7\uff1a\u516b\u5e74\u7ea7         \u2502',
    '\u2502 \u6458\u8981\uff1a\u2026\u2026             \u2502',
    '\u2502 \u9644\u4ef6\uff1axxx.docx       \u2502',
    '\u2502                     \u2502',
    '\u2502 [\u4fee\u6539\u8d44\u6599] [\u786e\u8ba4\u4fdd\u5b58] \u2502',
    '\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518',
])

heading(doc, '状态展示', level=3)
table(doc, ['状态', '移动端表现'],
    [['空对话', '欢迎卡片：\u201c上传课件或输入文字开始备课\u201d'],
     ['上传中', '附件卡片进度'],
     ['AI 思考中', '动态点 + 文案\u201cAI 正在分析内容\u201d'],
     ['轮询中', '卡片状态\u201c生成中\u201d'],
     ['普通聊天', '文本气泡'],
     ['结构化结果', 'AI 预览卡'],
     ['编辑中', '预览卡变为表单态'],
     ['保存中', '确认按钮 loading'],
     ['已保存', '\u201c已保存\u201d徽章'],
     ['保存失败', '红色错误提示 + 重试'],
     ['AI 超时', '超时提示卡片'],
     ['冲突', '覆盖 / 新建版本选项']],
    col_widths=[3, 9])

heading(doc, '移动端交互要点', level=3)
body(doc, '输入框固定底部；+ 弹出附件菜单；模式选择横向滑动；长按消息可复制；预览卡点击展开全部内容；确认保存使用底部确认栏；修改资料使用半屏表单弹层。')

heading(doc, '重要限制', level=3)
body(doc, 'AI 聊天输入包含 text、file_ids、mode 等字段，文件上传走现有 AI 上传能力。如果 Web 端没有真实实现语音或图片识别，小程序中不能新增语音/图片业务，只能在 + 菜单中保留\u201c文件上传\u201d能力，图片/语音不启用或作为后续扩展隐藏。')

# --- 2.4 空间首页 ---
page_break(doc)
heading(doc, '2.4 空间首页', level=2)
body(doc, '设计目的：Web 端顶部导航中有四个内容库。小程序底部只有 4 个 Tab，因此将四个内容库统一收纳到\u201c空间\u201dTab。这不是新增业务，而是移动端导航压缩。')

heading(doc, '页面布局', level=3)
flow(doc, ['Header：空间', 'Content：四个内容入口卡片 / 最近内容 / 筛选入口', 'Tab：底部导航'])

heading(doc, '四个入口', level=3)
flow(doc, ['个人备课    集体备课    计划与总结    教学反思'])

heading(doc, '组件结构', level=3)
table(doc, ['组件', '说明'],
    [['SpaceModuleGrid', '四个内容库入口'],
     ['RecentContentList', '最近内容'],
     ['StatusSummaryCard', '内容统计'],
     ['SearchBar', '全局搜索入口']],
    col_widths=[4, 8])

# --- 2.5 个人备课列表页 ---
heading(doc, '2.5 个人备课列表页', level=2)
body(doc, 'Web 端复用 ContentLibraryPage（表格+分页+抽屉）。移动端改为卡片列表，不改变数据和操作。')
heading(doc, '页面布局', level=3)
flow(doc, ['Header：个人备课 + 新建入口', 'Filter：搜索 / 学期 / 周次', 'Content：备课卡片列表'])
heading(doc, '卡片结构', level=3)
flow(doc, ['标题', '类型标签：个人备课', '学科 / 年级', '日期', '摘要', '附件标识', '状态标签', '操作：查看 / 编辑 / AI优化 / 删除'])
heading(doc, '移动端交互', level=3)
body(doc, '卡片点击进入详情；左滑显示\u201c编辑/删除\u201d；长按打开更多操作；分页改为上拉加载更多。')

# --- 2.6 集体备课列表页 ---
heading(doc, '2.6 集体备课列表页', level=2)
heading(doc, '卡片结构', level=3)
flow(doc, ['标题/主题', '类型标签：集体备课', '学科/年级', '日期', '摘要', '评论数', '附件标识', '操作：查看/编辑/评论/AI优化'])
heading(doc, '移动端交互', level=3)
body(doc, '点击卡片进入详情；评论入口放在详情页底部；左滑显示编辑和删除；上拉加载更多。')

# --- 2.7 计划与总结列表页 ---
heading(doc, '2.7 计划与总结列表页', level=2)
body(doc, '计划与总结属于统一内容类型 plan_summary，通过 plan_type 字段区分。')
heading(doc, '页面布局', level=3)
flow(doc, ['Header：计划与总结', 'Segment：全部 / 计划 / 总结', 'Filter：学期 / 周次 / 搜索', 'Content：卡片列表'])
heading(doc, '卡片结构', level=3)
flow(doc, ['标题', '标签：计划 / 总结', '学期', '日期', '摘要', '状态', '操作：查看/编辑/AI优化/导出'])
heading(doc, '移动端交互', level=3)
body(doc, '顶部 Segment 横向切换计划/总结；卡片点击进入详情；上拉加载更多。')

# --- 2.8 教学反思列表页 ---
heading(doc, '2.8 教学反思列表页', level=2)
body(doc, '教学反思需要关联已有备课（linkedContentId），详情中应展示被反思的备课信息。')
heading(doc, '卡片结构', level=3)
flow(doc, ['标题', '标签：教学反思', '关联备课', '日期', '反思摘要', '状态', '操作：查看/编辑/AI优化/导出/查看关联备课'])
heading(doc, '移动端交互', level=3)
body(doc, '点击\u201c关联备课\u201d进入备课详情；卡片点击进入反思详情；左滑编辑/删除。')

# --- 2.9 教师个人空间 ---
page_break(doc)
heading(doc, '2.9 教师个人空间', level=2)
body(doc, '展示单个教师的全部内容，支持统计筛选、内容类型切换。')
heading(doc, '页面布局', level=3)
flow(doc, ['Header：教师信息', 'Content：教师资料卡 → 内容统计卡 → 类型筛选 → 内容列表'])
heading(doc, '数据展示', level=3)
flow(doc, ['教师头像', '姓名', '部门/分组', '个人备课数', '集体备课数', '计划总结数', '教学反思数', '总数'])
heading(doc, '移动端交互', level=3)
body(doc, '统计卡横向滑动；内容类型 Tab 切换；卡片列表上拉加载。')

# --- 2.10 我的页面 ---
heading(doc, '2.10 我的页面', level=2)
body(doc, '\u201c我的\u201d不是新增业务，而是承接 Web 顶部用户区、个人信息、改密、退出登录，以及管理员后台入口。')
heading(doc, '页面布局', level=3)
flow(doc, ['Header：当前教师信息', 'Content：我的资料 / 修改密码 / 我的内容统计 / 管理后台入口（管理员可见） / 退出登录', 'Tab：底部导航'])
heading(doc, '管理后台入口说明', level=3)
body(doc, 'Web 端存在 /admin/* 系列后台页面。小程序教师端不建议把复杂后台全部塞进主 Tab，但管理员身份登录时，\u201c我的\u201d页需要保留\u201c管理后台\u201d入口，可打开管理后台 H5 或进入后续小程序后台分包，避免功能丢失。')

# --- 2.11 表单详情页 ---
heading(doc, '2.11 表单详情页', level=2)
body(doc, 'Web 端内容库使用表格+详情抽屉；小程序端改为独立详情页。')
heading(doc, '页面布局', level=3)
flow(doc, ['Header：标题 + 状态', 'Tabs：基本信息 / 内容正文 / 附件 / 评论或记录', 'Content：分块卡片', 'BottomAction：AI优化 / 编辑 / 导出 / 更多'])
heading(doc, '通用详情页结构', level=3)
flow(doc, [
    '标题区：标题、类型标签、状态标签、创建/更新时间',
    '基本信息：学科、年级、类型、关联内容',
    '内容正文：AI生成模块、正文内容、摘要',
    '附件：文件名、预览状态、下载入口',
    '评论/操作记录：评论列表、操作日志',
])
heading(doc, '底部操作栏', level=3)
flow(doc, ['AI优化    编辑    导出    更多'])
heading(doc, '移动端交互', level=3)
body(doc, '内容分块折叠/展开；附件点击预览；底部固定操作栏；更多操作用 ActionSheet。')

page_break(doc)

# ============ 3. 组件复用方案 ============
heading(doc, '3. 组件复用方案', level=1)

heading(doc, '3.1 基础组件', level=2)
table(doc, ['组件', '用途'],
    [['AppPage', '页面安全区容器'],
     ['AppNavBar', '顶部导航栏'],
     ['BottomTabBar', '底部导航'],
     ['SearchFilterBar', '搜索 + 筛选'],
     ['StatusTag', '状态标签'],
     ['TypeTag', '内容类型标签'],
     ['EmptyState', '空状态'],
     ['LoadingSkeleton', '加载骨架'],
     ['ConfirmModal', '确认弹窗'],
     ['ActionSheet', '更多操作']],
    col_widths=[4, 8])

heading(doc, '3.2 内容类组件', level=2)
table(doc, ['组件', '用途'],
    [['ContentCard', '个人备课/集体备课/计划总结/教学反思通用卡片'],
     ['ContentDetailBlock', '详情页内容分块'],
     ['AttachmentCard', '附件展示'],
     ['CommentList', '评论列表'],
     ['TeacherMiniCard', '首页教师卡片'],
     ['TeacherStatsCard', '教师统计卡']],
    col_widths=[4, 8])

heading(doc, '3.3 AI 类组件', level=2)
table(doc, ['组件', '用途'],
    [['ChatInputBar', 'AI 输入框'],
     ['AttachmentMenu', '附件菜单'],
     ['ModeSelector', 'AI 模式选择'],
     ['UserBubble', '用户消息'],
     ['AIBubble', 'AI 消息'],
     ['AIPreviewCard', 'AI 结构化预览卡'],
     ['AIPreviewEditSheet', 'AI 预览编辑弹层'],
     ['AISavingStatus', '保存状态'],
     ['AIUndoCard', '撤销提示']],
    col_widths=[4, 8])

heading(doc, '3.4 内容库复用逻辑', level=2)
body(doc, 'Web 端四个内容库复用 ContentLibraryPage，小程序端也应复用一个 ContentListPage：')
flow(doc, [
    'ContentListPage(kind)',
    '  kind = personal_lesson   → 个人备课列表',
    '  kind = group_lesson      → 集体备课列表',
    '  kind = plan_summary      → 计划总结列表',
    '  kind = reflection        → 教学反思列表',
])
body(doc, '这样移动端和 Web 端保持同一套业务抽象，不改变接口和数据结构。')

page_break(doc)

# ============ 4. 页面路由结构 ============
heading(doc, '4. 页面路由结构（小程序）', level=1)

heading(doc, '4.1 完整路由表', level=2)
flow(doc, [
    '/pages/login/index                    登录页',
    '/pages/home/index                     首页（教师名录）',
    '/pages/workspace/index                AI 工作台',
    '/pages/workspace/conversations/index  会话列表',
    '/pages/workspace/preview-edit/index   AI 预览编辑',
    '',
    '/pages/space/index                    空间首页',
    '/pages/space/personal-lessons/index   个人备课列表',
    '/pages/space/group-lessons/index      集体备课列表',
    '/pages/space/plans/index              计划总结列表',
    '/pages/space/reflections/index        教学反思列表',
    '',
    '/pages/content/detail/index           内容详情页',
    '/pages/content/edit/index             内容编辑页',
    '/pages/content/preview-file/index     文件预览页',
    '',
    '/pages/teacher/detail/index           教师个人空间',
    '',
    '/pages/mine/index                     我的',
    '/pages/mine/profile/index             个人资料',
    '/pages/mine/change-password/index     修改密码',
    '/pages/mine/admin-entry/index         管理后台入口',
])

heading(doc, '4.2 TabBar 配置', level=2)
table(doc, ['页面路径', 'Tab 名称'],
    [['/pages/home/index', '首页'],
     ['/pages/workspace/index', 'AI工作台'],
     ['/pages/space/index', '空间'],
     ['/pages/mine/index', '我的']],
    col_widths=[6, 6])

page_break(doc)

# ============ 5. AI 工作流在小程序中的表现形式 ============
heading(doc, '5. AI 工作流在小程序中的表现形式', level=1)

heading(doc, '5.1 小程序端完整链路', level=2)
body(doc, '以下链路完全对应 Web 当前 AI 工作流，不新增任何接口或业务逻辑。')
flow(doc, [
    '用户输入文字 / 上传文件',
    '        \u2193',
    'POST /api/ai/upload（如有附件）',
    '        \u2193',
    'POST /api/ai/chat',
    '        \u2193',
    '返回 sessionId / messageId / processing',
    '        \u2193',
    '轮询 GET /api/ai/recognition/:messageId',
    '        \u2193',
    'AI 返回普通回复 或 结构化预览卡',
    '        \u2193',
    '用户修改资料 / 确认保存 / 取消',
    '        \u2193',
    'POST /api/ai/confirm',
    '        \u2193',
    '保存 content + 子表 + 附件',
    '        \u2193',
    '对应内容库展示',
])

heading(doc, '5.2 AI 预览卡移动端展示', level=2)
flow(doc, [
    '\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510',
    '\u2502 AI\u8bc6\u522b\u7ed3\u679c          \u2502',
    '\u2502 \u7c7b\u578b\uff1a\u4e2a\u4eba\u5907\u8bfe       \u2502',
    '\u2502 \u6807\u9898\uff1a\u300a\u8377\u5858\u6708\u8272\u300b   \u2502',
    '\u2502 \u5b66\u79d1\uff1a\u8bed\u6587           \u2502',
    '\u2502 \u5e74\u7ea7\uff1a\u516b\u5e74\u7ea7         \u2502',
    '\u2502 \u6458\u8981\uff1a\u2026\u2026             \u2502',
    '\u2502 \u9644\u4ef6\uff1axxx.docx       \u2502',
    '\u2502                     \u2502',
    '\u2502 [\u4fee\u6539\u8d44\u6599] [\u786e\u8ba4\u4fdd\u5b58] \u2502',
    '\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518',
])

heading(doc, '5.3 AI 预览卡编辑', level=2)
body(doc, '编辑方式不再放在聊天气泡里复杂内联，而是点击\u201c修改资料\u201d后打开底部弹层：')
flow(doc, [
    '底部弹层字段：',
    '  - 类型 (type)',
    '  - 标题 (title)',
    '  - 学科 (subject)',
    '  - 年级 (grade)',
    '  - 日期 (date)',
    '  - 关联备课 (linkedContentId)',
    '  - 模块内容 (modules)',
    '  - 保存修改',
])
body(doc, '字段仍然使用 Web 当前 AI 预览卡字段，不新增字段。')

heading(doc, '5.4 保存成功后的表现', level=2)
flow(doc, [
    'AI消息卡片显示：已保存',
    '下方出现：查看详情 / 撤销',
    '对应内容进入空间列表',
])
body(doc, '撤销仍然使用 Web 现有 /api/ai/undo/:messageId 逻辑（5分钟窗口），不新增业务。')

page_break(doc)

# ============ 6. 关键设计结论 ============
heading(doc, '6. 关键设计结论', level=1)

conclusions = [
    ('6.1 首页不是传统\u201c工作台首页\u201d', '因为当前 Web /home 的真实功能是教师名录，所以小程序首页仍应以\u201c教师名录\u201d为主，而不是重新设计成统计大屏。可以增加移动端视觉优化，但不能改变首页核心功能。'),
    ('6.2 AI 工作台是小程序核心页', 'Web /workspace 是三栏 AI 工作台，小程序端必须保留：会话、聊天、模式选择、文件上传、AI 预览卡、修改资料、确认保存、最近内容查看。只是把三栏压缩为：聊天主页面 + 会话抽屉 + 最近内容抽屉 + 详情页。'),
    ('6.3 四个内容库必须复用一套移动端列表组件', '因为 Web 端四个页面已经复用 ContentLibraryPage，小程序也应复用 ContentListPage(kind)，避免重复开发。'),
    ('6.4 \u201c空间\u201dTab 是移动端导航压缩，不是新增功能', '空间页只是把个人备课、集体备课、计划与总结、教学反思集中展示，方便移动端进入。'),
    ('6.5 详情页要从\u201c抽屉\u201d改成\u201c独立页\u201d', 'Web 端详情是抽屉，小程序端屏幕窄，不适合抽屉详情，统一改为独立详情页，但内容结构、字段、接口保持不变。'),
    ('6.6 管理后台不建议塞进教师端主流程', 'Web 系统已有管理后台，但小程序主 Tab 应服务教师日常使用。管理员功能可以放在\u201c我的 → 管理后台入口\u201d，避免影响普通教师体验，也不删除管理员功能。'),
]

for title, desc in conclusions:
    heading(doc, title, level=2)
    body(doc, desc)

page_break(doc)

# ============ 7. 给 ZCode 的开发落地要求 ============
heading(doc, '7. 给 ZCode 的开发落地要求', level=1)
body(doc, '后续让 ZCode 开发时，必须强调以下 10 条规则：')

rules = [
    '不要重新设计业务流程',
    '不要新增接口',
    '不要新增字段',
    '不要改变 AI 工作流',
    '不要把 Web 表格直接搬到手机',
    '内容库统一做卡片化列表',
    '详情抽屉统一改为详情页',
    'AI 工作台优先保证\u201c输入、上传、轮询、预览卡、确认保存\u201d完整闭环',
    '图片/语音能力如果 Web 没有真实实现，不要在本期开放',
    '管理后台只做入口映射，不要影响教师端小程序主体验',
]

for i, rule in enumerate(rules, 1):
    body(doc, f'{i}. {rule}')

# ============ 保存 ============
os.makedirs(OUTPUT_DIR, exist_ok=True)
doc.save(OUTPUT_FILE)
print(f'\u2705 文档已生成: {OUTPUT_FILE}')
print(f'   文件大小: {os.path.getsize(OUTPUT_FILE):,} bytes')
