#!/usr/bin/env python3
"""
AI 教学系统 — 完整前端系统说明书 生成器
生成 Word 文档 (.docx)
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime
import os

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs')
OUTPUT_FILE = os.path.join(OUTPUT_DIR, '系统结构说明书.docx')

# ============ 样式常量 ============
FONT_CN_HEADING = 'SimHei'
FONT_CN_BODY = 'SimSun'
FONT_EN = 'Times New Roman'
COLOR_PRIMARY = RGBColor(0x10, 0x18, 0x20)
COLOR_BODY = RGBColor(0x18, 0x20, 0x30)
COLOR_ACCENT = RGBColor(0x1A, 0x56, 0xDB)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_BG_HEADER = '1A56DB'
COLOR_BG_ALT = 'F5F7FA'
COLOR_BORDER = 'D0D5DD'

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val","single")}" '
            f'w:sz="{val.get("sz","4")}" w:color="{val.get("color","D0D5DD")}" '
            f'w:space="0"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_run(paragraph, text, bold=False, size=Pt(10.5), font_cn=FONT_CN_BODY,
            font_en=FONT_EN, color=COLOR_BODY, italic=False):
    """添加格式化的文本运行"""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = font_en
    run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_cn}" w:ascii="{font_en}" '
        f'w:hAnsi="{font_en}"/>'
    )
    rPr.insert(0, rFonts)
    return run

def add_heading(doc, text, level=1):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_EN
        run.font.color.rgb = COLOR_PRIMARY
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_CN_HEADING}" '
            f'w:ascii="{FONT_EN}" w:hAnsi="{FONT_EN}"/>'
        )
        rPr.insert(0, rFonts)
    return h

def add_body(doc, text, indent=True):
    """添加正文段落"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    if indent:
        pf.first_line_indent = Pt(21)  # 2字符缩进
    add_run(p, text)
    return p

def add_code_block(doc, text):
    """添加代码块/等宽段落"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.left_indent = Cm(1)
    add_run(p, text, font_cn='Microsoft YaHei', size=Pt(9), color=RGBColor(0x30, 0x30, 0x30))
    return p

def add_table(doc, headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=Pt(9), font_cn=FONT_CN_HEADING, color=COLOR_WHITE)
        set_cell_shading(cell, COLOR_BG_HEADER)
    
    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            add_run(p, str(val), size=Pt(9))
            if r % 2 == 1:
                set_cell_shading(cell, COLOR_BG_ALT)
    
    # 列宽
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    
    doc.add_paragraph()  # 表后空行
    return table

def add_flow(doc, lines):
    """添加流程图（文本版）"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(1)
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run('\n')
        add_run(p, line, font_cn='Microsoft YaHei', size=Pt(8.5), color=RGBColor(0x40, 0x40, 0x40))
    doc.add_paragraph()
    return p

def add_page_break(doc):
    doc.add_page_break()

# ============ 创建文档 ============
doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.0)
section.right_margin = Cm(2.5)

# 设置默认字体
style = doc.styles['Normal']
style.font.name = FONT_EN
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN_BODY)

# ============ 封面 ============
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'AI 教学系统', bold=True, size=Pt(28), font_cn=FONT_CN_HEADING, color=COLOR_PRIMARY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '完整前端系统说明书', bold=True, size=Pt(20), font_cn=FONT_CN_HEADING, color=COLOR_ACCENT)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '技术栈：Next.js 15 + NestJS + BullMQ + PostgreSQL', size=Pt(11), color=RGBColor(0x50, 0x60, 0x70))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, f'生成日期：{datetime.date.today().strftime("%Y-%m-%d")}', size=Pt(11), color=RGBColor(0x50, 0x60, 0x70))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '本文档基于真实代码/数据库/路由/页面解析，无任何臆造', size=Pt(9), italic=True, color=RGBColor(0x80, 0x90, 0xA0))

add_page_break(doc)

# ============ 目录页 ============
add_heading(doc, '目录', level=1)
p = doc.add_paragraph()
add_run(p, '（请在 Word 中右键此处 → 更新域 → 更新整个目录）', size=Pt(9), italic=True, color=RGBColor(0x80, 0x90, 0xA0))

# 插入 TOC 域
paragraph = doc.add_paragraph()
run = paragraph.add_run()
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run._element.append(fldChar1)
run2 = paragraph.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
run2._element.append(instrText)
run3 = paragraph.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
run3._element.append(fldChar2)
run4 = paragraph.add_run('（请右键更新目录）')
fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run4._element.append(fldChar3)

add_page_break(doc)

# ============ 一、系统总览图 ============
add_heading(doc, '一、系统总览图', level=1)

add_body(doc, '本系统为 B/S 架构的 AI 教学备课平台，包含以下核心模块：')

add_flow(doc, [
    '┌─────────────────────────────────────────────────────────────────┐',
    '│                    前端 (Next.js 15 :8080)                       │',
    '│  /login  /home  /workspace  /personal-lessons  /group-lessons   │',
    '│  /plans  /reflections  /teacher/[id]  /admin/*                   │',
    '│  API Proxy → localhost:3000/api/*                                │',
    '└──────────────────────────┬──────────────────────────────────────┘',
    '                           │ HTTP (rewrite proxy)',
    '                           ▼',
    '┌─────────────────────────────────────────────────────────────────┐',
    '│                    API (NestJS :3000)                            │',
    '│  Auth / Teacher / School / Department / Content / Comments       │',
    '│  HomeGroup / AI Chat / AI Config / Conversation / Files          │',
    '│  BullMQ Queues: ai-recognition, file-preview, pdf-export         │',
    '└──────┬──────────────────┬──────────────────┬────────────────────┘',
    '       │                  │                  │',
    '       ▼                  ▼                  ▼',
    '┌──────────────┐  ┌──────────────┐  ┌──────────────────┐',
    '│  worker-ai   │  │worker-preview│  │worker-export/sched│',
    '│  (LLM调用)   │  │(LibreOffice) │  │ (PDF/stub)        │',
    '└──────┬───────┘  └──────────────┘  └──────────────────┘',
    '       │',
    '       ▼',
    '┌─────────────────────────────────────────────────────────────────┐',
    '│              DeepSeek / OpenAI / Qwen (LLM)                      │',
    '└─────────────────────────────────────────────────────────────────┘',
])

add_body(doc, '技术栈明细：前端 Next.js 15（App Router）；API NestJS 10（TypeORM + PostgreSQL）；队列 BullMQ（Redis）；Worker 服务 4 个（ai / preview / export / schedule）；LLM 适配器支持 DeepSeek、OpenAI、Qwen。')

add_page_break(doc)

# ============ 二、完整页面路由表 ============
add_heading(doc, '二、完整页面路由表', level=1)

add_body(doc, '以下列出所有前端页面路由及其基本信息：')

add_table(doc,
    ['#', 'URL 路径', '页面文件', '功能', '权限'],
    [
        ['1', '/', 'app/page.tsx', '根重定向（有 token → /home，无 → /login）', '公开'],
        ['2', '/login', 'app/login/page.tsx', '登录页（验证码+密码+首次改密）', '公开'],
        ['3', '/home', 'app/home/page.tsx', '教师名录（按部门/分组展示）', '登录'],
        ['4', '/workspace', 'app/workspace/page.tsx', 'AI 工作台（核心功能：AI 对话备课）', '登录'],
        ['5', '/personal-lessons', 'app/personal-lessons/page.tsx', '个人备课库（列表+详情+筛选）', '登录'],
        ['6', '/group-lessons', 'app/group-lessons/page.tsx', '集体备课库（列表+详情+筛选）', '登录'],
        ['7', '/plans', 'app/plans/page.tsx', '计划与总结库（列表+详情+筛选）', '登录'],
        ['8', '/reflections', 'app/reflections/page.tsx', '教学反思库（列表+详情+筛选）', '登录'],
        ['9', '/teacher/[id]', 'app/teacher/[id]/page.tsx', '教师个人空间（统计+内容列表/卡片）', '登录'],
        ['10', '/admin/teachers', 'app/admin/teachers/page.tsx', '管理员：教师管理（CRUD+导入）', '管理员'],
        ['11', '/admin/departments', 'app/admin/departments/page.tsx', '管理员：部门管理', '管理员'],
        ['12', '/admin/home-groups', 'app/admin/home-groups/page.tsx', '管理员：首页分组管理', '管理员'],
        ['13', '/admin/roles', 'app/admin/roles/page.tsx', '管理员：角色管理（Mock 数据）', '管理员'],
        ['14', '/admin/settings/school', '.../school/page.tsx', '管理员：学校信息设置', '管理员'],
        ['15', '/admin/settings/academic-years', '.../academic-years/page.tsx', '管理员：学年学期管理', '管理员'],
        ['16', '/admin/settings/ai-config', '.../ai-config/page.tsx', '管理员：AI 配置（Provider+Key+模型）', '管理员'],
        ['17', '/admin/settings/ai-logs', '.../ai-logs/page.tsx', '管理员：AI 操作日志', '管理员'],
    ],
    col_widths=[0.8, 3.2, 3.5, 4.5, 1.5]
)

add_page_break(doc)

# ============ 三、页面详细说明 ============
add_heading(doc, '三、页面详细说明', level=1)

# --- 3.1 登录页 ---
add_heading(doc, '3.1 登录页 /login', level=2)
add_body(doc, '功能：用户登录系统，包含验证码验证和首次登录强制改密。')

add_heading(doc, 'API 调用', level=3)
add_table(doc, ['API 端点', '方法', '说明'],
    [['GET /api/public/school', '公开', '获取学校品牌信息（Logo、名称、登录背景）'],
     ['POST /api/auth/login', '公开', '登录（mobile, password, captchaId, captchaCode）'],
     ['POST /api/auth/change-password', '需Token', '首次强制改密']],
    col_widths=[4, 1.5, 6])

add_heading(doc, '组件结构', level=3)
add_flow(doc, [
    'LoginPage',
    '├── 左侧品牌区',
    '│   ├── 学校 Logo 图片',
    '│   ├── 学校名称',
    '│   ├── 功能特性列表（4 项）',
    '│   └── 信任标识',
    '├── 右侧登录卡片',
    '│   ├── 手机号输入框',
    '│   ├── 密码输入框（可切换可见）',
    '│   ├── 验证码输入 + 验证码图片',
    '│   └── 登录按钮',
    '└── 强制改密弹窗（遮罩层）',
    '    ├── 新密码输入',
    '    ├── 确认密码输入',
    '    ├── 密码强度提示',
    '    └── 确认按钮',
])

add_heading(doc, '状态流', level=3)
add_flow(doc, [
    '未加载 → mounted(opacity过渡)',
    '├── 加载中（验证码skeleton / 登录中spinner）',
    '├── 错误（红色错误横幅 + 验证码自动刷新）',
    '├── mustChangePassword（弹窗遮罩）',
    '└── 成功 → 存token到localStorage → 跳转 /home',
])

add_heading(doc, '数据字段', level=3)
add_table(doc, ['字段', '类型', '来源', '必填'],
    [['mobile', 'string(11)', '用户输入', '是'],
     ['password', 'string(6-30)', '用户输入', '是'],
     ['captchaId', 'string', 'useCaptcha Hook', '是'],
     ['captchaCode', 'string', '用户输入', '是'],
     ['newPassword', 'string(8-32)', '用户输入（改密时）', '条件'],
     ['confirmPassword', 'string(8-32)', '用户输入（改密时）', '条件']],
    col_widths=[3, 2.5, 3.5, 1.5])

# --- 3.2 首页 ---
add_heading(doc, '3.2 教师名录首页 /home', level=2)
add_body(doc, '功能：展示全校教师列表，按部门/首页分组展示，支持搜索。')

add_table(doc, ['API 端点', '说明'],
    [['GET /api/home/groups', '获取首页可见的分组（含教师信息）'],
     ['GET /api/home/teachers?school_id=1', '获取所有教师'],
     ['GET /api/home/teachers-stats?school_id=1', '批量获取教师统计数']],
    col_widths=[5.5, 6.5])

add_flow(doc, [
    'HomePage',
    '├── TopNav（顶部导航栏）',
    '├── 页面标题 "教师名录"',
    '├── 搜索框',
    '├── 分组卡片列表（每个分组含教师网格）',
    '│   └── TeacherCard（头像 + 姓名 + 备课数 + 反思数）',
    '└── 空状态提示',
])

# --- 3.3 AI 工作台 ---
add_page_break(doc)
add_heading(doc, '3.3 AI 工作台 /workspace（核心页面）', level=2)
add_body(doc, '功能：AI 驱动的一站式备课工作台。支持文本/文件输入，AI 智能识别并生成备课内容预览卡，用户确认后自动落库。这是整个系统最核心的页面。')

add_heading(doc, '三栏布局组件结构', level=3)
add_flow(doc, [
    'WorkspacePage',
    '├── ConversationSidebar（左侧栏）',
    '│   ├── 新建会话按钮 / 搜索框',
    '│   └── 会话列表（按日期分组）',
    '├── 聊天区（中间 70%）',
    '│   ├── 消息列表',
    '│   │   ├── 用户消息气泡（蓝色，右对齐）',
    '│   │   └── AI 消息卡片（白色，左对齐，Sparkles图标）',
    '│   │       ├── 纯文本回复 / 结构化预览卡片',
    '│   │       │   ├── 类型标签 / 标题 / 日期',
    '│   │       │   ├── 学科/年级 / 来源/模块',
    '│   │       │   ├── "修改资料"按钮 → 编辑模式',
    '│   │       │   ├── "确认保存"按钮 / "取消"按钮',
    '│   │       │   └── 编辑模式：内联输入框',
    '│   ├── 底部输入区',
    '│   │   ├── + 弹出菜单（文件上传）',
    '│   │   ├── 文本输入框 / 模式选择 / 发送按钮',
    '├── 工作列表区（右侧 30%）',
    '│   ├── 筛选栏 / 手动添加按钮 / 备课列表',
    '│   └── 详情面板（元数据+摘要+附件+评论）',
    '└── 手动添加弹窗',
])

add_heading(doc, '状态流（完整链路）', level=3)
add_flow(doc, [
    '初始化 → 获取会话 → 获取消息',
    '├── 空对话（欢迎语："上传课件或输入文字开始备课"）',
    '├── 用户发送消息',
    '│   ├── 上传中 / AI思考中（动画点 + 步骤文字）',
    '│   ├── 轮询中（1.2秒/次，最多50次）',
    '│   ├── AI超时（"AI Worker 未响应"）',
    '│   ├── AI返回普通聊天（纯文本回复）',
    '│   └── AI返回备课预览卡',
    '│       ├── 只读预览 / 编辑中 / 保存中',
    '│       ├── 已保存（"✅ 已保存"徽章）',
    '│       ├── 冲突（覆盖/新建版本选项）',
    '│       └── 错误',
    '├── 查看备课详情（骨架屏/详情/删除确认/评论）',
    '└── 手动添加弹窗',
])

add_heading(doc, 'mode 枚举值', level=3)
add_table(doc, ['mode 值', '说明'],
    [['auto', '自动识别（默认，AI 判断意图）'],
     ['normal_chat', '普通聊天'],
     ['personal_lesson', '手动指定：个人备课'],
     ['group_lesson', '手动指定：集体备课'],
     ['semester_plan', '手动指定：学期计划'],
     ['semester_summary', '手动指定：学期总结'],
     ['teaching_reflection', '手动指定：教学反思']],
    col_widths=[4, 8])

add_heading(doc, 'ChatDto 发送字段', level=3)
add_table(doc, ['字段', '类型', '必填', '说明'],
    [['text', 'string', '否', '用户输入文本'],
     ['file_id', 'number/string', '否', '单个文件引用（旧）'],
     ['file_ids', 'number[]', '否', '多个文件引用'],
     ['scope', 'string', '否', "上下文（默认 'workspace'）"],
     ['mode', 'string', '否', "模式（默认 'auto'）"]],
    col_widths=[2.5, 2.5, 1, 6])

add_heading(doc, 'AI 识别结果字段', level=3)
add_table(doc, ['字段', '类型', '说明'],
    [['type', 'string', '识别类型（personal_lesson/reflection/group_lesson/plan_summary）'],
     ['title_candidate', 'string', 'AI建议标题'],
     ['subject', 'string', '学科'],
     ['grade', 'string', '年级'],
     ['summary', 'string', 'AI生成摘要'],
     ['confidence', 'number', '置信度'],
     ['modules', 'object', '模块化内容'],
     ['recognition_reasons', 'string[]', '识别依据'],
     ['showPreviewCard', 'boolean', '是否显示预览卡'],
     ['isBusinessScene', 'boolean', '是否业务场景'],
     ['need_user_confirm', 'boolean', '是否需要确认']],
    col_widths=[4, 2, 6])

add_heading(doc, '确认保存字段 ConfirmActionDto', level=3)
add_table(doc, ['字段', '类型', '必填', '说明'],
    [['messageId', 'number', '是', '关联消息ID'],
     ['type', 'string', '是', '内容类型'],
     ['title', 'string', '是', '标题'],
     ['subject', 'string', '否', '学科'],
     ['grade', 'string', '否', '年级'],
     ['linkedContentId', 'number', '否', '关联内容ID（反思用）'],
     ['extractedEntities', 'object', '否', '提取的实体'],
     ['fileIds', 'number[]', '否', '附件文件ID列表']],
    col_widths=[3, 2, 1.5, 5.5])

# --- 3.4-3.7 内容库页面 ---
add_page_break(doc)
add_heading(doc, '3.4-3.7 内容库页面', level=2)
add_body(doc, '四个页面（/personal-lessons, /group-lessons, /plans, /reflections）的结构和逻辑完全相同，均复用 ContentLibraryPage 组件，仅 kind 参数不同。')

add_table(doc, ['页面路由', 'kind 值', '页面标题'],
    [['/personal-lessons', 'personal_lesson', '个人备课'],
     ['/group-lessons', 'group_lesson', '集体备课'],
     ['/plans', 'plan_summary', '计划与总结'],
     ['/reflections', 'reflection', '教学反思']],
    col_widths=[4, 3, 5])

add_body(doc, '组件结构：ContentLibraryPage → TopNav → 页面头部（图标+标题+统计）→ FilterBar（学期/周次/搜索）→ 数据表格（序号/标题/类型/日期/操作）→ 分页 → 详情抽屉（Drawer+LessonDetailPanel）→ 删除确认弹窗。')
add_body(doc, '状态流：加载中（骨架屏4行）→ 无数据/"暂无XX" → 筛选无结果 → 查看详情（抽屉滑入）→ 删除确认 → 正常展示（表格+分页）。')

# --- 3.8 教师个人空间 ---
add_heading(doc, '3.8 教师个人空间 /teacher/[id]', level=2)
add_body(doc, '功能：展示单个教师的全部内容，支持统计筛选、列表/卡片模式切换。')
add_table(doc, ['API', '说明'],
    [['GET /api/home/teachers?school_id=1', '获取教师信息'],
     ['GET /api/teachers/{id}/contents?size=500', '获取教师全部内容'],
     ['GET /api/teachers/{id}/content-stats', '获取教师统计（personal_lesson, reflection, group_lesson, plan_summary, total）']],
    col_widths=[5.5, 6.5])

# --- 3.9-3.15 管理页面 ---
add_page_break(doc)
add_heading(doc, '3.9-3.15 管理后台页面', level=2)
add_body(doc, '管理后台统一使用 AdminShell 布局（280px 左侧导航 + 右侧内容区）。权限均为 JWT + @Roles(ADMIN)。')

add_table(doc, ['页面', '路由', '核心 API', '功能'],
    [['教师管理', '/admin/teachers', 'CRUD + 批量导入 + 查重 + 离职/恢复', '教师账号全生命周期管理'],
     ['部门管理', '/admin/departments', 'CRUD + 树形结构 + 教师分配', '组织架构管理'],
     ['首页分组', '/admin/home-groups', 'CRUD + 批量导入 + 教师绑定', '首页展示分组'],
     ['角色管理', '/admin/roles', 'Mock 数据', '角色权限管理（待接入）'],
     ['学校信息', '/admin/settings/school', 'GET/PUT /api/admin/school', '学校名称、Logo、登录背景'],
     ['学年学期', '/admin/settings/academic-years', 'GET/PUT /api/admin/school/settings', '学年学期配置'],
     ['AI 配置', '/admin/settings/ai-config', 'Provider CRUD + 测试 + Token统计', 'AI 服务配置管理'],
     ['AI 日志', '/admin/settings/ai-logs', 'GET /api/admin/ai-actions', 'AI 操作审计日志']],
    col_widths=[2, 2.5, 4, 3.5])

add_page_break(doc)

# ============ 四、AI 工作流完整链路 ============
add_heading(doc, '四、AI 工作流完整链路', level=1)

add_heading(doc, '4.1 端到端链路', level=2)
add_body(doc, '以下为用户从输入到内容落库的完整链路：')

add_flow(doc, [
    '1. 用户输入（文字 "帮我写一份三年级数学教案" / 上传文件 .docx）',
    '   │',
    '2. Frontend: POST /api/ai/chat',
    '   Payload: { text, file_id?, file_ids?, mode:\"auto\" }',
    '   │',
    '3. API (AIService.chat):',
    '   • 查找/创建 AISession (scope=\"workspace\")',
    '   • 保存 AIMessage (sender_type=\"teacher\")',
    '   • 读取文件内容 (txt→直接读, docx→解压XML, pdf→pdf-parse)',
    '   • 获取最近3条备课（用于反思关联）',
    '   • 入队 ai-recognition (classify-content)',
    '   • 返回 { sessionId, messageId, status:\"processing\" }',
    '   │',
    '4. worker-ai: classify-content 任务',
    '   Step 1: 去重检查（MD5 hash, 30s窗口）',
    '   Step 2: 用户限流（5次/分钟, 100次/天）',
    '   Step 3: 全局限流（60次/分钟, 5000次/天）',
    '   Step 4: 意图检测 INTENT_PROMPT → LLM',
    '       ├── CHAT/ASK → 普通聊天 → LLM(t=0.7, max=800) → 纯文本回复',
    '       ├── UPLOAD → 文件分析 → LLM UNIFIED_UPLOAD_PROMPT',
    '       └── CREATE/EDIT → 业务场景',
    '           Step 5: 任务检测 TASK_PROMPT → LLM',
    '           Step 6: 信息充分性检查（是否有年级+学科）',
    '           Step 7: 场景检测 detectScene()',
    '           Step 8: 并发保护',
    '           Step 9: LLM 场景化提示词 (t=0.1, json格式)',
    '           Step 10: 构建元数据 (cleanTitle, getTypeModules)',
    '   存入 Redis: ai_result:{messageId} (TTL 600s)',
    '   │',
    '5. Frontend: 轮询 GET /api/ai/recognition/{msgId}',
    '   间隔 1200ms, 最多 50 次 (60秒)',
    '   ├── 未找到 → 检查超时(120s) → timeout',
    '   └── 找到 → 创建 AI reply → 返回 { status:\"completed\", result:{...} }',
    '   │',
    '6. 用户操作预览卡片：修改资料 / 确认保存 / 取消',
    '   │',
    '7. POST /api/ai/confirm → ActionEngineService.execute [事务]',
    '   ① 版本冲突检测 ② 创建 Content ③ 创建子表记录',
    '   ④ 创建 LessonAttachment ⑤ 更新 AIRecognitionRecord',
    '   ⑥ 写入日志 ⑦ 存入 undo 快照 (TTL 300s)',
    '   ⑧ 创建 AI 成功回复 → 返回 { success, nl_reply, actionId }',
])

add_heading(doc, '4.2 Undo 撤销链路', level=2)
add_body(doc, 'POST /api/ai/undo/{msgId} → 检查 Redis undo:{msgId} (5分钟内有效) → 事务回滚（删除 Content + 子表 + 附件）→ 标记 ActionHistory 为 reverted → 创建 AI "已撤销"回复。')

add_heading(doc, '4.3 场景检测规则 detectScene', level=2)
add_table(doc, ['关键词', '场景', 'saveType'],
    [['教案、备课、课件、教学设计、课时', 'personal_lesson', 'personal_lesson'],
     ['集体备课、教研、集体', 'group_lesson', 'group_lesson'],
     ['学期计划、教学计划、工作计划', 'semester_plan', 'plan_summary'],
     ['学期总结、教学总结、工作总结', 'semester_summary', 'plan_summary'],
     ['教学反思、课后反思、反思', 'teaching_reflection', 'reflection'],
     ['其他', 'normal_chat', '-']],
    col_widths=[5, 3, 4])

add_page_break(doc)

# ============ 五、备课系统业务流 ============
add_heading(doc, '五、备课系统业务流', level=1)

# 5.1 个人备课
add_heading(doc, '5.1 个人备课流程', level=2)
add_flow(doc, [
    '触发方式：',
    '  ├── AI 工作台 → 输入文字（含"教案/备课"关键词）',
    '  ├── AI 工作台 → 上传课件 → AI分析',
    '  └── AI 工作台 → 手动选择 mode=personal_lesson',
    '',
    '流程：',
    '  1. 用户输入 → AI 意图检测 → CREATE → 任务检测 → Create Lesson',
    '  2. AI 生成结构化预览卡：type/personal_lesson, title_candidate, subject, grade, summary, modules',
    '  3. 用户可编辑预览卡内容，或直接确认',
    '  4. 确认保存：→ content 表 → personal_lesson 表 → lesson_attachment 表',
    '  5. 内容在 /personal-lessons 和 /teacher/[id] 可查看',
    '',
    '涉及表：content → personal_lesson → lesson_attachment → file_asset',
    '评论功能：非管理员限7天内，仅备课主人或管理员可评论',
])

# 5.2 集体备课
add_heading(doc, '5.2 集体备课流程', level=2)
add_flow(doc, [
    '触发方式：',
    '  ├── AI 工作台 → 输入文字（含"集体备课/教研"关键词）',
    '  └── AI 工作台 → 手动选择 mode=group_lesson',
    '',
    '流程：',
    '  1. 用户输入 → AI 场景检测 → group_lesson',
    '  2. AI 生成预览卡（topic, subject, grade, body_text）',
    '  3. 确认保存：→ content 表 → group_lesson 表',
    '  4. 内容在 /group-lessons 可查看',
    '',
    '涉及表：content → group_lesson → group_lesson_comment → file_asset',
    '评论功能：非管理员限3天内',
])

# 5.3 计划总结
add_heading(doc, '5.3 计划总结流程', level=2)
add_flow(doc, [
    '触发方式：',
    '  ├── AI 工作台 → 输入（含"学期计划/教学计划"→semester_plan）',
    '  ├── AI 工作台 → 输入（含"学期总结/教学总结"→semester_summary）',
    '  └── 手动选择 mode=semester_plan / semester_summary',
    '',
    '流程：',
    '  1. 用户输入 → AI 场景检测 → semester_plan/semester_summary',
    '  2. 确认保存：→ content 表(content_type=plan_summary) → plan_summary 表(plan_type, body_text)',
    '',
    '注意：plan_summary 是统一类型，通过 plan_type 字段区分',
    '  - semester_plan → plan_type="semester"',
    '  - semester_summary → plan_type="semester"',
])

# 5.4 教学反思
add_heading(doc, '5.4 教学反思流程', level=2)
add_flow(doc, [
    '触发方式：',
    '  ├── AI 工作台 → 输入（含"教学反思/课后反思"）',
    '  ├── 普通聊天中 AI 检测到反思内容关键词',
    '  └── 手动选择 mode=teaching_reflection',
    '',
    '流程：',
    '  1. 用户输入 → AI 场景检测 → teaching_reflection',
    '  2. AI 需要关联到某条已有备课（linkedContentId）',
    '  3. 确认保存：',
    '     → content 表(content_type=reflection)',
    '     → reflection 表(lesson_content_id, reflection_text)',
    '',
    '特殊逻辑：reflection 表有两个 content_id 关联：',
    '  - content_id: 反思自己的 content 记录',
    '  - lesson_content_id: 被反思的备课 content ID',
])

add_page_break(doc)

# ============ 六、表单结构汇总 ============
add_heading(doc, '六、表单结构汇总', level=1)

add_heading(doc, '6.1 登录表单', level=2)
add_table(doc, ['字段', '类型', '必填', 'AI生成', '可编辑', '验证规则'],
    [['mobile', 'tel input', '是', '否', '是', '11位手机号'],
     ['password', 'password', '是', '否', '是', '6-30位'],
     ['captchaCode', 'text', '是', '否', '是', '4位数字验证码']],
    col_widths=[2.5, 1.8, 1, 1, 1, 3])

add_heading(doc, '6.2 首次改密表单', level=2)
add_table(doc, ['字段', '类型', '必填', '验证规则'],
    [['newPassword', 'password', '是', '8-32位，大小写/数字/特殊字符至少两种'],
     ['confirmPassword', 'password', '是', '必须与newPassword一致']],
    col_widths=[3, 2, 2, 4])

add_heading(doc, '6.3 AI 聊天输入表单', level=2)
add_table(doc, ['字段', '类型', '必填', '说明'],
    [['text', 'textarea', '否', '用户输入文字'],
     ['file_ids', 'file[]', '否', '附件文件（可增删）'],
     ['mode', 'select', '否', 'auto/normal_chat/personal_lesson/...']],
    col_widths=[2.5, 2, 1, 5.5])

add_heading(doc, '6.4 AI 预览卡编辑表单', level=2)
add_table(doc, ['字段', '类型', '必填', 'AI生成', '可编辑', '说明'],
    [['type', 'select', '是', '是', '是', '可通过下拉修改类型'],
     ['title', 'text', '是', '是', '是', 'AI建议但可修改'],
     ['subject', 'text', '否', '是', '是', '学科'],
     ['grade', 'text', '否', '是', '是', '年级'],
     ['date', 'date', '否', '否', '是', '日期'],
     ['modules', 'object', '否', '是', '是（复杂）', '教学模块'],
     ['linkedContentId', 'select', '否（反思必填）', '是', '是', '关联的备课']],
    col_widths=[2.5, 1.5, 1.5, 1.5, 1.5, 3.5])

add_heading(doc, '6.5 教师创建/编辑表单（管理员）', level=2)
add_table(doc, ['字段', '类型', '创建时', '编辑时', '说明'],
    [['name', 'text', '✓', '✓', '姓名'],
     ['mobile', 'tel', '✓', '✗', '手机号（唯一标识）'],
     ['password', 'password', '✓', '✗', '密码'],
     ['gender', 'select', '✓', '✓', '男/女/其他'],
     ['department_ids', 'multi-select', '✓', '✓', '多部门'],
     ['role', 'multi-select', '✓', '✓', 'teacher/admin'],
     ['employee_no', 'text', '✓', '✓', '工号']],
    col_widths=[3, 2, 1.5, 1.5, 3])

add_heading(doc, '6.6 AI Provider 配置表单', level=2)
add_table(doc, ['字段', '类型', '必填', '说明'],
    [['name', 'text', '否', '配置名称'],
     ['provider_type', 'select', '是', 'deepseek/openai/qwen/custom'],
     ['api_key', 'password', '是', 'AES加密存储'],
     ['base_url', 'text', '否', 'API地址（按类型有默认值）'],
     ['default_model', 'text', '否', '默认模型'],
     ['models', 'text[]', '否', '可用模型列表'],
     ['enabled', 'boolean', '否', '是否启用'],
     ['is_active', 'boolean', '否', '是否当前使用']],
    col_widths=[3, 2, 1.5, 5.5])

add_page_break(doc)

# ============ 七、页面跳转关系图 ============
add_heading(doc, '七、页面跳转关系图', level=1)

add_flow(doc, [
    '                        ┌──────────────┐',
    '                        │      /       │',
    '                        │  (根重定向)   │',
    '                        └───┬──────┬───┘',
    '                 有token    │      │  无token',
    '                            ▼      ▼',
    '              ┌─────────┐    ┌──────────┐',
    '              │  /home   │    │  /login  │',
    '              │ 教师名录  │    │  登录页   │',
    '              └──┬───┬───┘    └────┬─────┘',
    '                 │   │             │ 登录成功',
    '                 │   └─────────────┘',
    '                 │',
    '    ┌────────────┼────────────┬──────────────┐',
    '    ▼            ▼            ▼              ▼',
    '┌─────────┐ ┌──────────┐ ┌───────┐ ┌──────────────┐',
    '│/workspace│ │/personal-│ │/group-│ │   /plans     │',
    '│AI工作台  │ │ lessons  │ │lessons│ │ 计划与总结    │',
    '└────┬─────┘ │ 个人备课  │ │集体备课│ └──────────────┘',
    '     │       └────┬─────┘ └───┬───┘',
    '     │            │           │          ┌──────────────┐',
    '     │            │           │          │ /reflections │',
    '     │            │           │          │  教学反思     │',
    '     │            │           │          └──────────────┘',
    '     │            ▼           ▼',
    '     │     ┌──────────────────────┐',
    '     │     │   ContentLibraryPage  │ ← 共用组件',
    '     │     │  (表格 + 抽屉详情)     │',
    '     │     └──────────────────────┘',
    '     │',
    '     ├──────────────────────────────────┐',
    '     ▼                                  ▼',
    '┌──────────────┐              ┌──────────────────┐',
    '│/teacher/[id] │              │   /admin/*       │',
    '│ 教师个人空间  │              │   管理后台        │',
    '└──────────────┘              └──────────────────┘',
])

add_body(doc, 'TopNav 导航链接：首页→/home | AI工作台→/workspace | 集体备课→/group-lessons | 个人备课→/personal-lessons | 计划与总结→/plans | 教学反思→/reflections')
add_body(doc, 'AdminShell 侧边栏：教师管理→/admin/teachers | 部门管理→/admin/departments | 首页分组→/admin/home-groups | 角色管理→/admin/roles | 学校信息→/admin/settings/school | 学年学期→/admin/settings/academic-years | AI配置→/admin/settings/ai-config | AI日志→/admin/settings/ai-logs')

add_page_break(doc)

# ============ 八、数据库实体一览 ============
add_heading(doc, '八、数据库实体一览', level=1)
add_body(doc, 'ORM: TypeORM | 数据库: PostgreSQL | 共 31 张表 | 使用 TypeORM Migration 管理（17 个迁移文件）')

add_table(doc, ['#', '表名', '实体类', '说明'],
    [['1', 'school', 'School', '学校信息（名称/Logo/背景/设置）'],
     ['2', 'department', 'Department', '组织架构（树形 parent_id）'],
     ['3', 'teacher', 'Teacher', '教师（登录账号/角色/多部门）'],
     ['4', 'teacher_status_history', 'TeacherStatusHistory', '教师状态变更审计'],
     ['5', 'content', 'Content', '统一内容表（单表继承）'],
     ['6', 'personal_lesson', 'PersonalLesson', '个人备课详情'],
     ['7', 'group_lesson', 'GroupLesson', '集体备课详情'],
     ['8', 'plan_summary', 'PlanSummary', '计划总结详情'],
     ['9', 'reflection', 'Reflection', '教学反思详情'],
     ['10', 'lesson_attachment', 'LessonAttachment', '内容-文件关联'],
     ['11', 'file_asset', 'FileAsset', '文件存储（带预览状态）'],
     ['12', 'preview_file', 'PreviewFile', '预览文件记录'],
     ['13', 'personal_lesson_comment', 'PersonalLessonComment', '个人备课评论'],
     ['14', 'group_lesson_comment', 'GroupLessonComment', '集体备课评论'],
     ['15', 'ai_session', 'AISession', 'AI 会话'],
     ['16', 'ai_message', 'AIMessage', 'AI 消息'],
     ['17', 'ai_recognition_record', 'AIRecognitionRecord', 'AI 识别记录'],
     ['18', 'ai_decision_log', 'AIDecisionLog', 'AI 决策日志'],
     ['19', 'ai_providers', 'AiProvider', 'AI Provider 注册表'],
     ['20', 'ai_configs', 'AiConfig', 'AI Provider 配置（密钥加密）'],
     ['21', 'ai_call_logs', 'AiCallLog', 'AI 调用日志（Token统计）'],
     ['22', 'ai_workspace', 'AiWorkspace', 'AI 工作区持久状态'],
     ['23', 'conversation', 'Conversation', 'AI 对话列表'],
     ['24', 'action_history', 'ActionHistory', 'AI 操作审计'],
     ['25', 'home_groups', 'HomeGroup', '首页分组'],
     ['26', 'home_group_teachers', 'HomeGroupTeacher', '分组-教师关联'],
     ['27', 'export_task', 'ExportTask', 'PDF 导出任务'],
     ['28', 'operation_log', 'OperationLog', '操作日志'],
     ['29', 'login_log', 'LoginLog', '登录日志'],
     ['30', 'deleted_record', 'DeletedRecord', '删除记录追踪'],
     ['31', 'system_config', 'SystemConfig', '系统键值配置']],
    col_widths=[0.8, 3.5, 2.8, 5])

add_heading(doc, '核心关系', level=2)
add_flow(doc, [
    'content (content_type: personal_lesson|group_lesson|reflection|plan_summary)',
    '  ├── 1:1 → personal_lesson (subject, grade, chapter, lesson_no, body_text)',
    '  ├── 1:1 → group_lesson (topic, grade, subject, body_text)',
    '  ├── 1:1 → reflection (lesson_content_id → content.id, reflection_text)',
    '  ├── 1:1 → plan_summary (plan_type, body_text)',
    '  └── 1:N → lesson_attachment → file_asset',
])

add_page_break(doc)

# ============ 九、API 端点完整清单 ============
add_heading(doc, '九、API 端点完整清单', level=1)
add_body(doc, '全局前缀 /api，响应格式 { code, message, data, requestId }，JWT Bearer Token 认证，角色分 teacher/admin。')

# 9.1 认证
add_heading(doc, '9.1 认证模块 /api/auth', level=2)
add_table(doc, ['方法', '路径', '权限', '说明'],
    [['GET', '/api/auth/captcha', '公开', '获取数学验证码（SVG）'],
     ['POST', '/api/auth/login', '公开', '登录'],
     ['POST', '/api/auth/refresh', '公开', '刷新 Token'],
     ['GET', '/api/auth/me', 'JWT', '获取当前用户信息'],
     ['POST', '/api/auth/logout', 'JWT', '登出'],
     ['POST', '/api/auth/change-password', 'JWT', '修改密码'],
     ['POST', '/api/auth/:teacherId/reset-password', 'Admin', '管理员重置密码']],
    col_widths=[1.2, 4.5, 1.2, 5])

# 9.2 学校
add_heading(doc, '9.2 学校模块', level=2)
add_table(doc, ['方法', '路径', '权限', '说明'],
    [['GET', '/api/public/school', '公开', '获取学校品牌信息'],
     ['GET', '/api/admin/school', 'JWT', '获取学校信息'],
     ['PUT', '/api/admin/school', 'Admin', '更新学校信息'],
     ['GET', '/api/admin/school/settings', 'JWT', '获取学年设置'],
     ['PUT', '/api/admin/school/settings', 'Admin', '更新学年设置']],
    col_widths=[1.2, 4.5, 1.2, 5])

# 9.3 教师
add_heading(doc, '9.3 教师模块', level=2)
add_table(doc, ['方法', '路径', '权限', '说明'],
    [['PUT', '/api/teacher/profile', 'JWT', '更新个人信息'],
     ['POST', '/api/teacher/change-password', 'JWT', '个人改密'],
     ['POST', '/api/teacher/avatar', 'JWT', '上传头像'],
     ['GET', '/api/admin/teachers', 'Admin', '教师列表'],
     ['POST', '/api/admin/teachers', 'Admin', '创建教师'],
     ['PUT', '/api/admin/teachers/:id', 'Admin', '更新教师'],
     ['POST', '/api/admin/teachers/:id/reset-password', 'Admin', '重置密码'],
     ['POST', '/api/admin/teachers/:id/disable', 'Admin', '停用'],
     ['POST', '/api/admin/teachers/:id/resign', 'Admin', '设为离职'],
     ['POST', '/api/admin/teachers/:id/restore', 'Admin', '恢复在职'],
     ['POST', '/api/admin/teachers/batch-import', 'Admin', '批量导入'],
     ['POST', '/api/admin/teachers/check-duplicates', 'Admin', '导入查重'],
     ['GET', '/api/home/teachers', 'JWT', '教师公开列表'],
     ['GET', '/api/home/teachers-stats', 'JWT', '批量教师统计']],
    col_widths=[1.2, 4.8, 1.2, 4.8])

# 9.4 内容
add_heading(doc, '9.4 内容模块', level=2)
add_table(doc, ['方法', '路径', '权限', '说明'],
    [['GET', '/api/teachers/:id/contents', 'JWT', '教师内容列表'],
     ['GET', '/api/teachers/:id/content-stats', 'JWT', '教师内容统计'],
     ['GET', '/api/contents/:id', 'JWT', '内容详情'],
     ['DELETE', '/api/contents/:id', 'JWT', '删除内容']],
    col_widths=[1.2, 4.5, 1.2, 5])

# 9.5 AI
add_heading(doc, '9.5 AI 模块', level=2)
add_table(doc, ['方法', '路径', '权限', '说明'],
    [['POST', '/api/ai/chat', 'JWT', '发送聊天消息'],
     ['GET', '/api/ai/recognition/:messageId', 'JWT', '轮询识别结果'],
     ['POST', '/api/ai/confirm/dry-run', 'JWT', '保存预览'],
     ['POST', '/api/ai/confirm', 'JWT', '确认保存'],
     ['POST', '/api/ai/undo/:messageId', 'JWT', '撤销操作（5分钟窗口）'],
     ['POST', '/api/ai/upload', 'JWT', '上传文件'],
     ['GET', '/api/ai/session', 'JWT', '获取活跃会话'],
     ['GET', '/api/ai/session/:id/messages', 'JWT', '获取会话消息'],
     ['GET', '/api/ai/chat-quota', 'JWT', '获取每日配额']],
    col_widths=[1.2, 4.8, 1.2, 4.8])

# 9.6 评论
add_heading(doc, '9.6 评论模块', level=2)
add_table(doc, ['方法', '路径', '权限', '说明'],
    [['GET', '/api/group-lessons/:id/comments', 'JWT', '集体备课评论列表'],
     ['POST', '/api/group-lessons/:id/comments', 'JWT', '添加评论（3天限制）'],
     ['DELETE', '/api/group-lessons/comments/:id', 'JWT', '删除评论'],
     ['GET', '/api/personal-lessons/:id/comments', 'JWT', '个人备课评论列表'],
     ['POST', '/api/personal-lessons/:id/comments', 'JWT', '添加评论（7天限制）'],
     ['DELETE', '/api/personal-lessons/comments/:id', 'JWT', '删除评论']],
    col_widths=[1.2, 4.8, 1.2, 4.8])

# 9.7 AI 配置
add_heading(doc, '9.7 AI 配置管理 /api/admin/ai-configs', level=2)
add_table(doc, ['方法', '路径', '权限', '说明'],
    [['GET', '/api/admin/ai-configs', 'Admin', 'Provider列表'],
     ['POST', '/api/admin/ai-configs', 'Admin', '创建Provider'],
     ['PUT', '/api/admin/ai-configs/:id', 'Admin', '更新Provider'],
     ['DELETE', '/api/admin/ai-configs/:id', 'Admin', '删除Provider'],
     ['POST', '/api/admin/ai-configs/:id/test', 'Admin', '测试连接'],
     ['POST', '/api/admin/ai-configs/:id/activate', 'Admin', '设为当前'],
     ['POST', '/api/admin/ai-configs/:id/enable', 'Admin', '启用'],
     ['POST', '/api/admin/ai-configs/:id/disable', 'Admin', '停用'],
     ['GET', '/api/admin/ai-configs/stats', 'Admin', 'Token统计'],
     ['GET', '/api/admin/ai-configs/stats/users', 'Admin', '用户统计'],
     ['GET', '/api/admin/ai-configs/active/internal', '公开', 'Worker获取解密配置'],
     ['POST', '/api/admin/ai-configs/log-call', '公开', 'Worker记录调用'],
     ['GET', '/api/admin/ai-actions', 'Admin', 'AI操作日志']],
    col_widths=[1.2, 4.8, 1.2, 4.8])

add_page_break(doc)

# ============ 十、统计汇总 ============
add_heading(doc, '十、统计汇总', level=1)

add_table(doc, ['项目', '数量', '说明'],
    [['前端页面', '17 个', '含1个重定向路由'],
     ['API 端点组', '17 组', '含约 70+ 个具体端点'],
     ['数据库表', '31 张', '含 17 个迁移文件'],
     ['Worker 服务', '4 个', 'ai / preview / export / schedule'],
     ['AI 场景', '6 种', 'normal_chat/personal_lesson/group_lesson/semester_plan/semester_summary/teaching_reflection'],
     ['备课类型', '4 种', '个人备课/集体备课/计划总结/教学反思'],
     ['用户角色', '2 种', 'teacher / admin'],
     ['ORM 框架', 'TypeORM', 'PostgreSQL 数据库'],
     ['前端框架', 'Next.js 15', 'App Router 模式'],
     ['API 框架', 'NestJS 10', 'JWT + RolesGuard'],
     ['消息队列', 'BullMQ', '基于 Redis'],
     ['LLM 支持', '3 种', 'DeepSeek / OpenAI / Qwen']],
    col_widths=[3, 2, 7])

# ============ 保存 ============
os.makedirs(OUTPUT_DIR, exist_ok=True)
doc.save(OUTPUT_FILE)
print(f'✅ 文档已生成: {OUTPUT_FILE}')
print(f'   文件大小: {os.path.getsize(OUTPUT_FILE):,} bytes')
